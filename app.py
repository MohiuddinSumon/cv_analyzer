import json
import logging
import os
import re
from typing import Any, Dict, List, Optional
import requests
from bs4 import BeautifulSoup
from datetime import datetime

# LLM integration
import anthropic
import cv2
import docx
import numpy as np
import openai

# Document processing libraries
import PyPDF2
import pytesseract

# Web interface
import streamlit as st
from pdf2image import convert_from_path
from streamlit_chat import message
from tenacity import retry, stop_after_attempt, wait_exponential

# Set up logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)
from dotenv import load_dotenv


class CVProcessor:
    """Handles the document processing and information extraction from CVs."""

    def __init__(self, ocr_engine="tesseract"):
        """Initialize the CV processor with the specified OCR engine."""
        self.ocr_engine = ocr_engine
        # Configure OCR
        if ocr_engine == "tesseract":
            # Try to auto-detect Tesseract path or use environment variable
            try:
                pytesseract.get_tesseract_version()
                logger.info("Tesseract found in system PATH")
            except:
                # Check common installation paths
                common_paths = [
                    r"/usr/bin/tesseract",
                    r"/usr/local/bin/tesseract",
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    os.environ.get("TESSERACT_PATH"),
                ]

                for path in common_paths:
                    if path and os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        logger.info(f"Tesseract found at: {path}")
                        break
                else:
                    logger.warning("Tesseract not found. OCR may fail.")

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using multiple methods."""
        text = ""

        # Method 1: Direct text extraction
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                logger.info(
                    f"Direct extraction from {file_path}: {len(text)} characters"
                )
        except Exception as e:
            logger.error(f"Error in direct PDF extraction: {str(e)}")

        # Method 2: If direct extraction failed or yielded minimal text, try OCR
        if len(text.strip()) < 50:  # Lower threshold
            try:
                logger.info(f"Using OCR for {file_path}")
                ocr_text = self._extract_text_with_ocr(file_path)

                # If OCR yields more text, use it
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    logger.info(f"OCR produced better results: {len(text)} characters")
            except Exception as e:
                logger.error(f"OCR processing failed: {str(e)}")
                # Keep the original text if OCR fails
        # print(f"Extracted texts from PDF {text}")
        return text

    def _extract_text_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR with better error handling."""
        text = ""
        try:
            # Check if poppler is available for pdf2image
            try:
                # Convert PDF to images
                images = convert_from_path(file_path)
            except Exception as e:
                logger.error(f"Error converting PDF to images: {str(e)}")
                logger.info(
                    "Ensure poppler is installed. On Ubuntu: sudo apt-get install poppler-utils"
                )
                return ""

            for i, image in enumerate(images):
                # Convert PIL image to OpenCV format
                open_cv_image = np.array(image)

                # Handle both RGB and grayscale images
                if len(open_cv_image.shape) == 3:
                    open_cv_image = open_cv_image[
                        :, :, ::-1
                    ].copy()  # RGB to BGR conversion
                    gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
                else:
                    gray = open_cv_image.copy()

                # Try different preprocessing techniques
                # 1. Simple thresholding
                thresh = cv2.threshold(
                    gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )[1]

                # Perform OCR
                page_text = pytesseract.image_to_string(thresh)

                # If results are poor, try adaptive thresholding
                if len(page_text.strip()) < 50:
                    logger.info(f"Trying adaptive threshold for page {i+1}")
                    adaptive_thresh = cv2.adaptiveThreshold(
                        gray,
                        255,
                        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                        cv2.THRESH_BINARY,
                        11,
                        2,
                    )
                    page_text = pytesseract.image_to_string(adaptive_thresh)

                text += page_text + "\n"
                logger.info(f"Extracted {len(page_text)} characters from page {i+1}")

            return text
        except Exception as e:
            logger.error(f"OCR processing failed for {file_path}: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files with better error handling."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"Extracted {len(text)} characters from DOCX {file_path}")
            # print(f"Extracted texts from DOCX {text}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""

    def process_document(self, file_path: str) -> str:
        """Process a document and extract its text content."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""

        file_extension = os.path.splitext(file_path)[1].lower()

        logger.info(f"Processing document: {file_path} with extension {file_extension}")

        if file_extension == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_extension in [".docx", ".doc"]:
            if file_extension == ".doc":
                logger.warning(
                    ".doc format requires conversion to .docx for best results"
                )
            return self.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return ""


class CVAnalyzer:
    """Uses LLM to extract structured information from CV text."""

    def __init__(self, llm_provider="gemini", api_key=None):
        """Initialize the CV analyzer with the specified LLM provider."""
        self.llm_provider = llm_provider

        if llm_provider == "anthropic":
            self.client = anthropic.Anthropic(
                api_key=api_key or os.getenv("ANTHROPIC_API_KEY")
            )
            self.model = "claude-3-7-sonnet-20250219"
        elif llm_provider == "openai":
            self.client = openai.OpenAI(api_key=api_key or os.getenv("OPENAI_API_KEY"))
            self.model = "gpt-4"
        elif llm_provider == "gemini":
            # Import Google's Gemini library
            import google.generativeai as genai

            # Configure the Gemini API with your API key
            genai.configure(api_key=api_key or os.getenv("GOOGLE_API_KEY"))

            # Store the genai module for later use
            self.genai = genai
            self.model = "gemini-2.0-flash"
        else:
            raise ValueError(f"Unsupported LLM provider: {llm_provider}")

    @retry(
        stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=60)
    )
    def extract_cv_information(self, cv_text: str) -> Dict[str, Any]:
        """Extract structured information from CV text using LLM."""
        prompt = f"""
        You are a CV analysis expert. Extract the following structured information from the CV text below:
        
        1. Personal Information (name, email, phone, location, portfolio, github_url, linkedin_url)
        2. Education History [(institution, degree, field, graduation_date)]
        3. Work Experience [(company, role, duration, [responsibilities], [achievements])]
        4. Skills ([technical], [soft], [languages])
        5. Projects [(name, description, technologies)]
        6. Certifications [(name, issuer, date)]
        
        Return the information in JSON format with these exact categories. If any information is not found, include the category with an empty value or array. Follow the formatting of the data. 
        
        CV Text:
        {cv_text}
        """

        try:
            if self.llm_provider == "anthropic":
                response = self.client.messages.create(
                    model=self.model,
                    max_tokens=4000,
                    messages=[{"role": "user", "content": prompt}],
                )
                result = response.content[0].text
            elif self.llm_provider == "openai":
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "You are a CV analysis expert."},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=4000,
                )
                result = response.choices[0].message.content
            elif self.llm_provider == "gemini":
                # Generate content using Gemini
                generation_config = {
                    "temperature": 0.2,
                    "top_p": 0.95,
                    "top_k": 0,
                    "max_output_tokens": 4000,
                }

                # Initialize the model
                model = self.genai.GenerativeModel(
                    model_name=self.model, generation_config=generation_config
                )

                # Send the prompt directly as a string, not as a structured message
                response = model.generate_content(prompt)

                # Extract the text from the response
                result = response.text
                print(f"LLM Gemini Response: {result}")

            # Extract JSON from the response
            json_match = re.search(r"```json\n([\s\S]*?)\n```", result)
            if json_match:
                json_str = json_match.group(1)
            else:
                json_str = result

            # Clean up any non-JSON content
            json_str = re.sub(r"^[^{]*", "", json_str)
            json_str = re.sub(r"[^}]*$", "", json_str)

            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Error extracting CV information: {str(e)}")
            # Return a basic empty structure on error
            return {
                "personal_information": {},
                "education_history": [],
                "work_experience": [],
                "skills": {},
                "projects": [],
                "certifications": [],
            }


class CVDatabase:
    """Stores and manages extracted CV information."""

    def __init__(self, storage_path="cv_database.json"):
        """Initialize the CV database with the specified storage path."""
        self.storage_path = storage_path
        self.cvs = {}
        self.load_database()

    def load_database(self):
        """Load the CV database from storage."""
        try:
            if os.path.exists(self.storage_path):
                with open(self.storage_path, "r") as file:
                    self.cvs = json.load(file)
        except Exception as e:
            logger.error(f"Error loading CV database: {str(e)}")
            self.cvs = {}

    def save_database(self):
        """Save the CV database to storage."""
        try:
            with open(self.storage_path, "w") as file:
                json.dump(self.cvs, file, indent=2)
        except Exception as e:
            logger.error(f"Error saving CV database: {str(e)}")

    def add_cv(self, cv_id: str, cv_data: Dict[str, Any]):
        """Add a CV to the database."""
        self.cvs[cv_id] = cv_data
        self.save_database()

    def get_cv(self, cv_id: str) -> Optional[Dict[str, Any]]:
        """Get a CV from the database by ID."""
        return self.cvs.get(cv_id)

    def get_all_cvs(self) -> Dict[str, Dict[str, Any]]:
        """Get all CVs from the database."""
        return self.cvs

    def search_cvs(self, query: str) -> List[str]:
        """Basic search functionality for CVs."""
        query = query.lower()
        results = []

        for cv_id, cv_data in self.cvs.items():
            # Search in skills
            skills = cv_data.get("skills", {})
            all_skills = []
            for skill_category in skills.values():
                if isinstance(skill_category, list):
                    all_skills.extend(skill_category)

            if any(query in skill.lower() for skill in all_skills):
                results.append(cv_id)
                continue

            # Search in work experience
            work_exp = cv_data.get("work_experience", [])
            for exp in work_exp:
                if (
                    query in exp.get("company", "").lower()
                    or query in exp.get("role", "").lower()
                ):
                    results.append(cv_id)
                    break

            # Search in education
            education = cv_data.get("education_history", [])
            for edu in education:
                if (
                    query in edu.get("institution", "").lower()
                    or query in edu.get("field", "").lower()
                ):
                    results.append(cv_id)
                    break

        return results


class ProfileDatabase:
    """Stores and manages user profile information."""

    def __init__(self, storage_path="profile_database.json"):
        """Initialize the profile database with the specified storage path."""
        self.storage_path = storage_path
        self.profile_data = {}
        self.load_database()

    def load_database(self):
        """Load the profile database from storage."""
        try:
            if os.path.exists(self.storage_path):
                with open(self.storage_path, "r") as file:
                    self.profile_data = json.load(file)
        except Exception as e:
            logger.error(f"Error loading profile database: {str(e)}")
            self.profile_data = {}

    def save_database(self):
        """Save the profile database to storage."""
        try:
            with open(self.storage_path, "w") as file:
                json.dump(self.profile_data, file, indent=2)
        except Exception as e:
            logger.error(f"Error saving profile database: {str(e)}")

    def get_profile(self) -> Dict[str, Any]:
        """Get the user's profile data."""
        return self.profile_data

    def update_profile(self, profile_data: Dict[str, Any]):
        """Update the user's profile data."""
        self.profile_data = profile_data
        self.save_database()


class CVQueryEngine:
    """Handles natural language queries about the CV database."""

    def __init__(self, cv_database: CVDatabase, llm_provider="gemini", api_key=None):
        """Initialize the query engine with the CV database and LLM provider."""
        self.cv_database = cv_database
        self.llm_provider = llm_provider
        self.conversation_context = []

        if llm_provider == "anthropic":
            self.client = anthropic.Anthropic(
                api_key=api_key or os.getenv("ANTHROPIC_API_KEY")
            )
            self.model = "claude-3-7-sonnet-20250219"
        elif llm_provider == "openai":
            self.client = openai.OpenAI(api_key=api_key or os.getenv("OPENAI_API_KEY"))
            self.model = "gpt-4"
        elif llm_provider == "gemini":
            # Import Google's Gemini library
            import google.generativeai as genai

            # Configure the Gemini API with your API key
            genai.configure(api_key=api_key or os.getenv("GOOGLE_API_KEY"))

            # Store the genai module for later use
            self.genai = genai
            self.model = "gemini-2.0-flash"
        else:
            raise ValueError(f"Unsupported LLM provider: {llm_provider}")

    def add_to_context(self, role: str, content: str):
        """Add a message to the conversation context."""
        self.conversation_context.append({"role": role, "content": content})
        # Keep only the last 10 messages for context window management
        if len(self.conversation_context) > 10:
            self.conversation_context = self.conversation_context[-10:]

    @retry(
        stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=60)
    )
    def process_query(self, query: str) -> str:
        """Process a natural language query about the CVs."""
        all_cvs = self.cv_database.get_all_cvs()

        # Create a compact representation of CVs for context
        cv_summary = {}
        for cv_id, cv_data in all_cvs.items():
            # Handle both formats: "personal_information" and "Personal Information"
            personal_info = cv_data.get(
                "personal_information", cv_data.get("Personal Information", {})
            )

            name = personal_info.get("name", f"Candidate {cv_id}")
            email = personal_info.get("email", "No email provided")
            phone = personal_info.get("phone", "No phone number provided")
            location = personal_info.get("location", "Unknown location")

            # Extract key skills - handle both formats
            skills = []
            skills_data = cv_data.get("skills", cv_data.get("Skills", {}))

            for skill_category, skill_list in skills_data.items():
                if isinstance(skill_list, list):
                    skills.extend(
                        skill_list[:15]
                    )  # Increased limit to 15 skills per category

            # Extract latest work experience - handle both formats
            latest_work = "No work experience"
            work_exp = cv_data.get(
                "work_experience", cv_data.get("Work Experience", [])
            )

            if work_exp:
                latest = work_exp[0]
                role = latest.get("role", "Role")
                company = latest.get("company", "Company")
                latest_work = f"{role} at {company}"

            # Extract latest education - handle both formats
            latest_edu = "No education information"
            education = cv_data.get(
                "education_history", cv_data.get("Education History", [])
            )

            if education:
                latest = education[0]
                degree = latest.get("degree", "Degree")
                field = latest.get("field", "Field")
                institution = latest.get("institution", "Institution")
                latest_edu = f"{degree} in {field} from {institution}"

            cv_summary[cv_id] = {
                "name": name,
                "email": email,
                "phone": phone,
                "location": location,
                "latest_work": latest_work,
                "latest_education": latest_edu,
                "top_skills": skills[:30],  # Increased overall limit to 30 skills
            }

        # print(f"CV summary: {json.dumps(cv_summary, indent=2)}")

        # Prepare the conversation for the LLM
        system_prompt = f"""
        You are a CV analysis assistant. You have access to {len(all_cvs)} CV profiles with the following summary information:
        
        {json.dumps(cv_summary, indent=2)}
        
        Provide helpful, accurate responses to queries about these candidates. If you need more specific information that's not in the summary, indicate that you may need to look into the detailed CV.
        """

        # Add the query to the context
        self.add_to_context("user", query)

        try:
            if self.llm_provider == "anthropic":
                messages = [{"role": "system", "content": system_prompt}]
                messages.extend(self.conversation_context)

                response = self.client.messages.create(
                    model=self.model, max_tokens=2000, messages=messages
                )
                result = response.content[0].text
            elif self.llm_provider == "openai":
                messages = [{"role": "system", "content": system_prompt}]
                messages.extend(self.conversation_context)

                response = self.client.chat.completions.create(
                    model=self.model, messages=messages, max_tokens=2000
                )
                result = response.choices[0].message.content
            elif self.llm_provider == "gemini":
                # Create a properly formatted conversation history for Gemini
                gemini_messages = []

                # Add system prompt as a user message at the beginning
                gemini_messages.append({"role": "user", "parts": [system_prompt]})

                # Add a placeholder response from the model to acknowledge the system prompt
                gemini_messages.append(
                    {
                        "role": "model",
                        "parts": [
                            "I understand the CV profiles and will provide helpful responses."
                        ],
                    }
                )

                # Add the conversation context
                for message in self.conversation_context:
                    role = "user" if message["role"] == "user" else "model"
                    gemini_messages.append(
                        {"role": role, "parts": [message["content"]]}
                    )

                # Generate content using Gemini
                generation_config = {
                    "temperature": 0.2,
                    "top_p": 0.95,
                    "top_k": 0,
                    "max_output_tokens": 2000,
                }

                # Initialize the model
                model = self.genai.GenerativeModel(
                    model_name=self.model, generation_config=generation_config
                )

                # Start a chat and send the full history
                chat = model.start_chat(
                    history=gemini_messages[:-1]
                )  # Exclude the last user message
                response = chat.send_message(
                    gemini_messages[-1]["parts"][0]
                )  # Send the last user message

                result = response.text

            # Add the response to the context
            self.add_to_context("assistant", result)
            return result
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            return f"I encountered an error while processing your query. Please try again or rephrase your question. Error details: {str(e)}"

    def get_detailed_cv_info(self, cv_id: str) -> Optional[Dict[str, Any]]:
        """Get detailed information for a specific CV."""
        return self.cv_database.get_cv(cv_id)


class CVAnalysisSystem:
    """Main class that coordinates the CV analysis system components."""

    def __init__(
        self,
        ocr_engine="tesseract",
        llm_provider="gemini",
        api_key=None,
        storage_path="cv_database.json",
        profile_storage_path="profile_database.json",
    ):
        """Initialize the CV analysis system."""
        self.cv_processor = CVProcessor(ocr_engine=ocr_engine)
        self.cv_analyzer = CVAnalyzer(llm_provider=llm_provider, api_key=api_key)
        self.cv_database = CVDatabase(storage_path=storage_path)
        self.profile_database = ProfileDatabase(storage_path=profile_storage_path)
        self.query_engine = CVQueryEngine(
            self.cv_database, llm_provider=llm_provider, api_key=api_key
        )

    def process_cv_file(self, file_path: str) -> str:
        """Process a CV file and store its structured information."""
        try:
            # Generate a unique ID for the CV - remove 'temp_' prefix if present
            base_filename = os.path.basename(file_path)
            if base_filename.startswith("temp_"):
                base_filename = base_filename[5:]  # Remove the 'temp_' prefix

            cv_id = base_filename.split(".")[0]

            # Check if this CV has already been processed
            if self.cv_database.get_cv(cv_id):
                logger.info(f"CV {cv_id} already exists in the database")
                return cv_id

            # Extract text from the document
            logger.info(f"Extracting text from {file_path}")
            cv_text = self.cv_processor.process_document(file_path)

            if not cv_text:
                logger.error(f"Failed to extract text from {file_path}")
                return None

            # Extract structured information using LLM
            logger.info(f"Analyzing CV content for {file_path}")
            cv_data = self.cv_analyzer.extract_cv_information(cv_text)

            # Store the CV data
            logger.info(f"Storing CV data for {file_path}")
            self.cv_database.add_cv(cv_id, cv_data)

            return cv_id
        except Exception as e:
            logger.error(f"Error processing CV file {file_path}: {str(e)}")
            return None

    def process_cv_directory(self, directory_path: str) -> List[str]:
        """Process all CV files in a directory."""
        processed_ids = []

        for filename in os.listdir(directory_path):
            if filename.endswith((".pdf", ".docx", ".doc")):
                file_path = os.path.join(directory_path, filename)
                cv_id = self.process_cv_file(file_path)
                if cv_id:
                    processed_ids.append(cv_id)

        return processed_ids

    def query_cvs(self, query: str) -> str:
        """Query the CV database using natural language."""
        return self.query_engine.process_query(query)

    def get_cv_details(self, cv_id: str) -> Optional[Dict[str, Any]]:
        """Get detailed information for a specific CV."""
        return self.cv_database.get_cv(cv_id)


def fetch_job_description(url: str) -> str:
    """Fetch job description from a URL."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Get text content
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        logger.error(f"Error fetching job description: {str(e)}")
        return ""


# Streamlit web interface
def run_web_interface(llm_provider="gemini", api_key=None):
    """Run the Streamlit web interface for the CV analysis system."""
    st.set_page_config(page_title="CV Analysis System", layout="wide")

    # Initialize the CV analysis system
    if "cv_system" not in st.session_state:
        st.session_state.cv_system = CVAnalysisSystem(
            llm_provider=llm_provider, api_key=api_key
        )

    # Initialize chat history
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # Initialize selected CV state
    if "selected_cv" not in st.session_state:
        st.session_state.selected_cv = None

    # Initialize job analysis history
    if "job_analysis_history" not in st.session_state:
        st.session_state.job_analysis_history = []

    # Initialize profile data from persistent storage
    if "profile_data" not in st.session_state:
        st.session_state.profile_data = st.session_state.cv_system.profile_database.get_profile()

    st.title("CV Analysis System")

    # Sidebar for file upload and processing
    with st.sidebar:
        st.header("Upload and Process CVs")

        uploaded_files = st.file_uploader(
            "Upload CV files", type=["pdf", "docx", "doc"], accept_multiple_files=True
        )

        if uploaded_files:
            process_button = st.button("Process CVs")

            if process_button:
                with st.spinner("Processing CV files..."):
                    for uploaded_file in uploaded_files:
                        # Save the uploaded file temporarily
                        temp_file_path = f"temp_{uploaded_file.name}"
                        with open(temp_file_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())

                        # Process the CV
                        cv_id = st.session_state.cv_system.process_cv_file(
                            temp_file_path
                        )

                        if cv_id:
                            st.success(f"Processed CV: {uploaded_file.name}")
                        else:
                            st.error(f"Failed to process CV: {uploaded_file.name}")

                        # Remove the temporary file
                        os.remove(temp_file_path)

        # Display the list of processed CVs with improved formatting
        st.header("Processed CVs")
        all_cvs = st.session_state.cv_system.cv_database.get_all_cvs()

        if all_cvs:
            for cv_id, cv_data in all_cvs.items():
                # Extract name and email from personal information
                personal_info = cv_data.get("Personal Information", {})
                name = personal_info.get("name", "")
                email = personal_info.get("email", "")

                # Format the display text
                if name or email:
                    display_text = ""
                    if name:
                        display_text += f"{name}"
                    if email:
                        display_text += f" ({email})"
                else:
                    display_text = f"Candidate {cv_id}"

                # Check if this is the profile resume

                is_profile = cv_id == st.session_state.get("profile_data", {}).get("resume_id", None)
                
                # Create a clickable link for each CV with different styling for profile
                if is_profile:
                    st.markdown(f"<div style='background-color: #e6ffe6; padding: 10px; border-radius: 5px; margin: 5px 0;'>", unsafe_allow_html=True)
                    st.markdown("**My Profile Resume:**")
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        if st.button(display_text, key=f"btn_{cv_id}"):
                            st.session_state.selected_cv = cv_id
                    with col2:
                        if st.button("❌", key=f"delete_{cv_id}"):
                            # Delete from database
                            del st.session_state.cv_system.cv_database.cvs[cv_id]
                            st.session_state.cv_system.cv_database.save_database()
                            
                            # Clear profile data if this was the profile resume
                            if st.session_state.profile_data.get("resume_id") == cv_id:
                                st.session_state.profile_data["resume_id"] = None
                            
                            # Clear selected CV if this was selected
                            if st.session_state.selected_cv == cv_id:
                                st.session_state.selected_cv = None
                            
                            st.rerun()
                    st.markdown("</div>", unsafe_allow_html=True)
                else:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        if st.button(display_text, key=f"btn_{cv_id}"):
                            st.session_state.selected_cv = cv_id
                    with col2:
                        if st.button("❌", key=f"delete_{cv_id}"):
                            # Delete from database
                            del st.session_state.cv_system.cv_database.cvs[cv_id]
                            st.session_state.cv_system.cv_database.save_database()
                            
                            # Clear selected CV if this was selected
                            if st.session_state.selected_cv == cv_id:
                                st.session_state.selected_cv = None
                            
                            st.rerun()
        else:
            st.sidebar.write("No CVs processed yet")

    # Main area with tabs for chat and CV details
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Chat Assistant", "CV Details", "Profile", "Job Analysis", "Resume Editor"])

    with tab1:
        # Display chat history
        for i, (query, response) in enumerate(st.session_state.chat_history):
            message(query, is_user=True, key=f"user_{i}")
            message(response, key=f"assistant_{i}")

        # Chat input
        user_query = st.chat_input(
            "Ask a question about the Candidates:", key="user_query"
        )

        if user_query:
            # Save the query to prevent duplicate processing on rerun
            st.session_state.last_query = user_query

            with st.spinner("Processing your query..."):
                response = st.session_state.cv_system.query_cvs(user_query)
                st.session_state.chat_history.append((user_query, response))

            # Clear the input after sending
            st.rerun()

    with tab2:
        if st.session_state.selected_cv:
            cv_id = st.session_state.selected_cv
            cv_data = st.session_state.cv_system.get_cv_details(cv_id)

            if cv_data:
                st.header(f"CV Details: {cv_id}")

                # Display formatted JSON
                st.json(cv_data)

                # Add option to view original document if needed in the future
                st.write("Note: Original document viewer could be added here")
        else:
            st.info("Select a CV from the sidebar to view details")

    with tab3:
        st.header("My Profile")
        
        # Profile Resume Upload
        st.subheader("My Resume")
        profile_resume = st.file_uploader(
            "Upload your resume", type=["pdf", "docx", "doc"], key="profile_resume"
        )

        if profile_resume:
            with st.spinner("Processing your resume..."):
                # Save the uploaded file temporarily
                temp_file_path = f"temp_{profile_resume.name}"
                with open(temp_file_path, "wb") as f:
                    f.write(profile_resume.getbuffer())

                # Process the CV
                cv_id = st.session_state.cv_system.process_cv_file(temp_file_path)
                
                if cv_id:
                    st.session_state.profile_data["resume_id"] = cv_id
                    
                    # Get the processed CV data
                    profile_cv = st.session_state.cv_system.get_cv_details(cv_id)
                    
                    # Auto-populate links if found in the CV
                    personal_info = profile_cv.get("personal_information".replace('_', ' ').title(), {})
                    st.session_state.profile_data["portfolio_link"] = personal_info.get("portfolio", "")
                    st.session_state.profile_data["github_link"] = personal_info.get("github_url", "")
                    st.session_state.profile_data["linkedin_link"] = personal_info.get("linkedin_url", "")
                    
                    # Auto save to persistent storage
                    st.session_state.cv_system.profile_database.update_profile(st.session_state.profile_data)
                    st.success("Your resume has been processed and profile updated successfully!")
                else:
                    st.error("Failed to process your resume")

                # Remove the temporary file
                os.remove(temp_file_path)

        # Profile Links
        st.subheader("Professional Links")
        st.session_state.profile_data["portfolio_link"] = st.text_input(
            "Portfolio Link", value=st.session_state.profile_data.get("portfolio_link", "")
        )
        st.session_state.profile_data["github_link"] = st.text_input(
            "GitHub Link", value=st.session_state.profile_data.get("github_link", "")
        )
        st.session_state.profile_data["linkedin_link"] = st.text_input(
            "LinkedIn Link", value=st.session_state.profile_data.get("linkedin_link", "")
        )

        # Additional Information
        st.subheader("Additional Information")
        st.session_state.profile_data["additional_info"] = st.text_area(
            "Additional Information",
            value=st.session_state.profile_data.get("additional_info", ""),
            height=150
        )

        # Save Profile Button
        if st.button("Save Profile"):
            st.session_state.cv_system.profile_database.update_profile(st.session_state.profile_data)
            st.success("Profile saved successfully!")

        # Display analyzed profile information if available
        if st.session_state.profile_data.get("resume_id"):
            profile_cv = st.session_state.cv_system.get_cv_details(
                st.session_state.profile_data["resume_id"]
            )
            
            if profile_cv:
                st.subheader("Analyzed Profile Information")
                # Display the full CV data in JSON format
                st.json(profile_cv)
            else:
                st.warning("Unable to load profile data. Please try uploading your resume again.")

    with tab4:
        st.header("Job Analysis")
        
        # Job Description Input
        st.subheader("Job Description")
        
        # URL input and fetch button
        col1, col2 = st.columns([3, 1])
        with col1:
            job_url = st.text_input("Job Posting URL (optional)")
        with col2:
            fetch_button = st.button("Fetch Details", disabled=not job_url)
            
        if fetch_button:
            with st.spinner("Fetching job description..."):
                fetched_description = fetch_job_description(job_url)
                if fetched_description:
                    st.session_state.job_description = fetched_description
                    st.success("Successfully fetched job description!")
                else:
                    st.error("Failed to fetch job description from the URL. Please paste it manually.")
        
        # Job description text area
        job_description = st.text_area(
            "Paste Job Description",
            value=st.session_state.get("job_description", ""),
            height=200
        )
        
        if st.button("Analyze Job Match"):
            if not job_description:
                st.warning("Please provide a job description to analyze.")
            elif not st.session_state.profile_data.get("resume_id"):
                st.warning("Please upload your resume in the Profile tab first.")
            else:
                with st.spinner("Analyzing job match..."):
                    # Get the profile CV data
                    profile_cv = st.session_state.cv_system.get_cv_details(
                        st.session_state.profile_data["resume_id"]
                    )
                    
                    # Create analysis prompt
                    analysis_prompt = f"""
                    Analyze the following CV against the job description and provide:
                    1. Match percentage and key strengths
                    2. Missing skills or experience
                    3. Specific recommendations for improvement
                    4. Suggested resume modifications
                    
                    CV Data:
                    {json.dumps(profile_cv, indent=2)}
                    
                    Job Description:
                    {job_description}
                    """
                    
                    # Get analysis from LLM
                    analysis = st.session_state.cv_system.query_engine.process_query(analysis_prompt)
                    
                    # Store analysis in history
                    analysis_entry = {
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "job_url": job_url,
                        "job_description": job_description,
                        "analysis": analysis
                    }
                    st.session_state.job_analysis_history.append(analysis_entry)
                    
                    # Display analysis
                    st.subheader("Analysis Results")
                    st.write(analysis)
        
        # Separate Cover Letter Generation Section
        if job_description:  # Only show if there's a job description
            st.subheader("Generate Cover Letter")
            company_name = st.text_input("Company Name")
            hiring_manager = st.text_input("Hiring Manager Name (optional)")
            
            if st.button("Generate Cover Letter", key="generate_cover_letter"):
                if not st.session_state.profile_data.get("resume_id"):
                    st.warning("Please upload your resume in the Profile tab first.")
                else:
                    with st.spinner("Generating cover letter..."):
                        # Get the profile CV data
                        profile_cv = st.session_state.cv_system.get_cv_details(
                            st.session_state.profile_data["resume_id"]
                        )
                        
                        # Get the latest analysis if available
                        latest_analysis = ""
                        if st.session_state.job_analysis_history:
                            latest_analysis = st.session_state.job_analysis_history[-1].get("analysis", "")
                        
                        # Create cover letter prompt
                        cover_letter_prompt = f"""
                        Generate a professional cover letter based on the following information:
                        
                        Job Description:
                        {job_description}
                        
                        Candidate's CV:
                        {json.dumps(profile_cv, indent=2)}
                        
                        Job Analysis Results:
                        {latest_analysis}
                        
                        Company Name: {company_name}
                        Hiring Manager: {hiring_manager if hiring_manager else 'Hiring Manager'}
                        
                        Please create a compelling cover letter that:
                        1. Addresses the specific job requirements
                        2. Highlights relevant experience and skills from the CV
                        3. Incorporates insights from the job analysis
                        4. Uses a professional tone and format
                        5. Is concise and impactful
                        """
                        
                        # Generate cover letter using LLM
                        cover_letter = st.session_state.cv_system.query_engine.process_query(cover_letter_prompt)
                        
                        # Store cover letter in the latest analysis entry
                        if st.session_state.job_analysis_history:
                            st.session_state.job_analysis_history[-1]["cover_letter"] = cover_letter
                        
                        # Display the generated cover letter
                        st.subheader("Generated Cover Letter")
                        st.write(cover_letter)
                        
                        # Add download button for the cover letter
                        st.download_button(
                            label="Download Cover Letter",
                            data=cover_letter,
                            file_name=f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain"
                        )

        # Display analysis history with cover letters
        if st.session_state.job_analysis_history:
            st.subheader("Previous Analyses and Cover Letters")
            for entry in reversed(st.session_state.job_analysis_history):
                with st.expander(f"Analysis from {entry['timestamp']}"):
                    if entry['job_url']:
                        st.write(f"**Job URL:** {entry['job_url']}")
                    st.write("**Job Description:**")
                    st.write(entry['job_description'])
                    st.write("**Analysis:**")
                    st.write(entry['analysis'])
                    
                    if entry.get('cover_letter'):
                        st.write("**Cover Letter:**")
                        st.write(entry['cover_letter'])
                    
                    # Download buttons for analysis and cover letter
                    col1, col2 = st.columns(2)
                    with col1:
                        analysis_text = f"""
                        Job URL: {entry['job_url']}
                        Timestamp: {entry['timestamp']}
                        
                        Job Description:
                        {entry['job_description']}
                        
                        Analysis:
                        {entry['analysis']}
                        """
                        st.download_button(
                            label="Download Analysis",
                            data=analysis_text,
                            file_name=f"job_analysis_{entry['timestamp'].replace(' ', '_')}.txt",
                            mime="text/plain"
                        )
                    
                    if entry.get('cover_letter'):
                        with col2:
                            st.download_button(
                                label="Download Cover Letter",
                                data=entry['cover_letter'],
                                file_name=f"cover_letter_{entry['timestamp'].replace(' ', '_')}.txt",
                                mime="text/plain",
                                key=f"dl_cover_letter_{entry['timestamp']}"
                            )

    with tab5:
        st.header("Resume Editor")
        
        if not st.session_state.profile_data.get("resume_id"):
            st.warning("Please upload your resume in the Profile tab first.")
        else:
            # Get the profile CV data
            profile_cv = st.session_state.cv_system.get_cv_details(
                st.session_state.profile_data["resume_id"]
            )
            
            if profile_cv:
                # Display editable sections
                st.subheader("Edit Resume Sections")
                
                # Personal Information
                st.write("Personal Information")
                personal_info = profile_cv.get("personal_information".replace('_', ' ').title(), {})
                edited_name = st.text_input("Name", value=personal_info.get("name", ""))
                edited_email = st.text_input("Email", value=personal_info.get("email", ""))
                edited_phone = st.text_input("Phone", value=personal_info.get("phone", ""))
                edited_location = st.text_input("Location", value=personal_info.get("location", ""))
                
                # Skills
                st.write("Skills")
                skills = profile_cv.get("skills".title(), {})
                technical_skills = st.text_area(
                    "Technical Skills",
                    value="\n".join(skills.get("technical", [])),
                    height=100
                )
                soft_skills = st.text_area(
                    "Soft Skills",
                    value="\n".join(skills.get("soft", [])),
                    height=100
                )
                
                # Work Experience
                st.write("Work Experience")
                work_exp = profile_cv.get("work_experience".replace('_', ' ').title(), [])
                for i, exp in enumerate(work_exp):
                    with st.expander(f"Experience {i+1}"):
                        exp["company"] = st.text_input("Company", value=exp.get("company", "N/A"), key=f"company_{i}")
                        exp["role"] = st.text_input("Role", value=exp.get("role", ""), key=f"role_{i}")
                        exp["duration"] = st.text_input("Duration", value=exp.get("duration", ""), key=f"duration_{i}")
                        # Convert responsibilities list to string for text area
                        responsibilities_text = "\n".join(exp.get("responsibilities", []))
                        edited_responsibilities = st.text_area(
                            "Responsibilities (one per line)",
                            value=responsibilities_text,
                            height=100,
                            key=f"responsibilities_{i}"
                        )
                        # Convert back to list, removing empty lines
                        exp["responsibilities"] = [r.strip() for r in edited_responsibilities.split("\n") if r.strip()]
                
                # Save Changes Button
                if st.button("Save Changes"):
                    # Update the CV data
                    updated_cv = {
                        "personal_information": {
                            "name": edited_name,
                            "email": edited_email,
                            "phone": edited_phone,
                            "location": edited_location
                        },
                        "skills": {
                            "technical": [s.strip() for s in technical_skills.split("\n") if s.strip()],
                            "soft": [s.strip() for s in soft_skills.split("\n") if s.strip()]
                        },
                        "work_experience": work_exp
                    }
                    
                    # Save to database
                    st.session_state.cv_system.cv_database.add_cv(
                        st.session_state.profile_data["resume_id"],
                        updated_cv
                    )
                    
                    st.success("Resume updated successfully!")
            else:
                st.warning("Unable to load profile data. Please try uploading your resume again.")


if __name__ == "__main__":

    load_dotenv()

    # Try to get Google API key first
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        print("GOOGLE_API_KEY not found. Checking for ANTHROPIC_API_KEY...")
        api_key = os.getenv("ANTHROPIC_API_KEY")
        llm_provider = "anthropic"
    else:
        llm_provider = "gemini"

    if not api_key:
        print(
            "Please set either GOOGLE_API_KEY or ANTHROPIC_API_KEY environment variable"
        )
        exit(1)

    print(f"Using {llm_provider} as the LLM provider")
    cv_system = CVAnalysisSystem(llm_provider=llm_provider, api_key=api_key)

    # Process a directory of CVs
    # cv_dir = "data/sample_cvs"
    # if os.path.exists(cv_dir):
    #     print(f"Processing CVs in {cv_dir}...")
    #     processed_ids = cv_system.process_cv_directory(cv_dir)
    #     print(f"Processed {len(processed_ids)} CVs")

    # Run the web interface
    run_web_interface(llm_provider=llm_provider, api_key=api_key)
